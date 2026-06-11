"""生成技术方案简要说明 Word 文档"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
    return h


def add_para(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(10.5)
    run.bold = bold
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "微软雅黑"
                run.font.size = Pt(9)
                run.bold = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run.font.size = Pt(9)
    return table


LQ = "\u201c"  # 左双引号 "
RQ = "\u201d"  # 右双引号 "

# ═══════════════════════════════════════════
# 标题
# ═══════════════════════════════════════════
title = doc.add_heading("MES 智能问数系统 \u2014 技术方案简要说明", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"

# ═══════════════════════════════════════════
# 1. 核心思路
# ═══════════════════════════════════════════
add_heading("一、核心思路", 1)

add_para(
    "不要用 AI 解决一个可以用工程确定性解决的问题。"
    "高频、有明确口径的查询走指标层（100% 可靠），"
    "长尾、复杂、自定义的查询才走 NL2SQL。"
)

add_para("方案不是替代现有系统，而是在其前面加一层指标路由层，形成双层架构：")

add_para(
    "用户查询 \u2192 术语归一化 \u2192 路由判断 \u2192 指标通道（查视图，100% 可靠）或 SQL 通道（LLM 生成，现有方案）"
)

# ═══════════════════════════════════════════
# 2. 双层架构
# ═══════════════════════════════════════════
add_heading("二、双层架构", 1)

add_table(
    ["维度", "指标通道", "SQL 通道"],
    [
        ["适用场景", "口径固定的高频查询", "长尾、复杂关联查询"],
        ["查询方式", "预定义 PostgreSQL 视图", "LLM 实时生成 SQL"],
        ["可靠性", "100%（不经过 AI）", "~80%"],
        ["响应速度", "毫秒级", "秒级"],
        ["口径权威性", "业务方书面确认后固化", "每次生成可能不同"],
        ["目标占比", "3 个月后 > 60%", "3 个月后 < 40%"],
    ],
)

add_para("")
add_para("路由判断逻辑：", bold=True)

add_table(
    ["状态", "含义", "处理方式"],
    [
        ["matched", "精确命中单个指标", "直接查视图，返回结果"],
        ["ambiguous", f"术语有歧义（如{LQ}合格率{RQ}多口径）", "返回追问，让用户选择"],
        ["multi_match", "命中多个指标", "分别查询视图，合并展示"],
        ["no_match", "未匹配任何指标", "走现有 NL2SQL 通道"],
    ],
)

# ═══════════════════════════════════════════
# 3. 术语归一化
# ═══════════════════════════════════════════
add_heading("三、术语归一化方案", 1)

add_para("核心设计：使用代码节点中的精确字典做别名映射，不用向量检索。", bold=True)
add_para(
    f"理由是：别名映射是精确查找问题（{LQ}良品率{RQ}\u2192M004），"
    f"不需要语义相似度来猜。知识库只负责{LQ}从 380 张表里找相关表{RQ}，不参与术语归一化。"
)

add_para("别名字典示例：")
add_table(
    ["用户说法", "映射指标", "说明"],
    [
        ["良品率、合格率、直通率、一次合格率", "M004 过站良品率", "同口径多别名"],
        ["首次通过率、FPY、首通率", "M005 首次通过率 FPY", "含英文缩写"],
        ["来料合格率、IQC合格率、进料检验合格率", "M007 IQC合格率", ""],
        ["产量、日产量、生产数量、完工数量", "M001 工单日产量", ""],
        ["在制工单数、在制数、WIP", "M003 在制工单数", ""],
        ["库存、现有库存、实时库存", "M011 实时库存", ""],
    ],
)

add_para("")
add_para("歧义术语特殊处理：")
add_para(
    f"{LQ}合格率{RQ}在不同部门至少有 4 种口径（过站良品率、IQC、IPQC、FQC），"
    "不能直接映射到一个指标上。这类术语进入歧义列表，"
    f"系统追问用户{LQ}你要查的是哪种合格率？{RQ}再执行。"
)

# ═══════════════════════════════════════════
# 4. 动态参数与 SQL 拼装
# ═══════════════════════════════════════════
add_heading("四、动态参数提取与 SQL 拼装", 1)

add_para(
    f"视图的 SQL 是固定的，但{LQ}今天{RQ}{LQ}A线{RQ}{LQ}P001{RQ}"
    "等过滤条件每次不同，需要从用户输入中提取参数，再拼接成完整的查询 SQL。"
)

add_para("参数提取规则：")
add_table(
    ["参数类型", "用户说法示例", "提取方式", "输出示例"],
    [
        ["今天/昨日", "今天/今日", "关键词 + datetime 计算", "date: 2026-06-05"],
        ["本周/本月", "本月/这个月", "关键词 + 日期计算", "start: 月初, end: 今天"],
        ["具体日期", "2024-01-15", "正则匹配日期格式", "date: 2024-01-15"],
        ["产线", "A线/SMT线", "正则匹配", "pdline: A线"],
        ["料号", "P001/SKU-001", "正则匹配大写字母+数字", "part_no: P001"],
        ["工序", "贴片/AOI/ICT", "工序关键词匹配", "process: 贴片"],
        ["供应商", "供应商XX", "正则提取", "supplier: XX"],
        ["无时间词", "良品率多少", "默认近 7 天", "start: 7天前, end: 今天"],
    ],
)

add_para("")
add_para("完整执行示例：", bold=True)
add_para(f"用户输入：{LQ}今天A线良品率多少{RQ}")
add_para(f"  \u2192 术语归一化：命中{LQ}良品率{RQ} \u2192 M004")
add_para("  \u2192 参数提取：{date: 2026-06-05, pdline: A线}")
add_para("  \u2192 SQL 拼装：SELECT ... FROM v_m004_yield_rate WHERE 统计日期='2026-06-05' AND 产线名称 ILIKE '%A线%'")
add_para("  \u2192 执行查询 \u2192 返回表格结果")

# ═══════════════════════════════════════════
# 5. 第一期指标清单
# ═══════════════════════════════════════════
add_heading("五、第一期指标清单（17 个视图）", 1)

add_para("生产类：", bold=True)
add_table(
    ["编号", "指标名称", "视图名", "确认状态"],
    [
        ["M001", "工单日产量", "v_m001_wo_daily_output", "待确认"],
        ["M002", "工单计划达成率", "v_m002_wo_achievement", "待确认"],
        ["M003", "在制工单数", "v_m003_wip_count", "已确认"],
        ["M004", "SN过站良品率", "v_m004_yield_rate", "已确认（核心指标）"],
        ["M005", "首次通过率 FPY", "v_m005_fpy", "已确认"],
        ["M006", "工单平均周期", "v_m006_wo_cycle_time", "待确认"],
    ],
)

add_para("")
add_para("质量类：", bold=True)
add_table(
    ["编号", "指标名称", "视图名", "确认状态"],
    [
        ["M007", "IQC来料合格率", "v_m007_iqc_rate", "已确认"],
        ["M008", "IPQC巡检合格率", "v_m008_ipqc_rate", "已确认"],
        ["M009", "FQC成品合格率", "v_m009_fqc_rate", "已确认"],
        ["M010", "TOP N不良统计", "v_m010_top_defects", "已确认"],
    ],
)

add_para("")
add_para("库存类：", bold=True)
add_table(
    ["编号", "指标名称", "视图名", "确认状态"],
    [
        ["M011", "实时库存", "v_m011_stock", "已确认"],
        ["M012", "当日领料量", "v_m012_daily_issue", "已确认"],
        ["M014", "采购到货准时率", "v_m014_po_ontime", "待确认"],
    ],
)

add_para("")
add_para("设备类：", bold=True)
add_table(
    ["编号", "指标名称", "视图名", "确认状态"],
    [
        ["M015", "设备月故障次数", "v_m015_equipment_failure", "已确认"],
        ["M016", "平均维修时长 MTTR", "v_m016_mttr", "待确认"],
        ["M017", "设备点检完成率", "v_m017_inspection_rate", "待确认"],
    ],
)

# ═══════════════════════════════════════════
# 6. 与现有系统集成
# ═══════════════════════════════════════════
add_heading("六、与现有系统集成", 1)

add_para(
    "保留现有：LangGraph 工作流、BFS 图扩展、pgvector 向量检索、"
    "Harness 数据飞轮、SQL 安全校验——全部保留，仅在 SQL 通道内使用。"
)

add_para("")
add_para("新增部分：")
add_table(
    ["组件", "说明"],
    [
        ["术语归一化代码节点", "Python 精确字典匹配 + 参数提取"],
        ["指标路由节点", "根据匹配状态分发到指标通道或 SQL 通道"],
        ["指标视图查询节点", "拼装参数化 SQL \u2192 HTTP 调用只读数据库"],
        ["歧义追问节点", "当术语有歧义时向用户返回追问消息"],
        ["17 个 PostgreSQL 视图", "部署在 MES 业务库只读 Schema"],
    ],
)

add_para("")
add_para("Dify 工作流示意：", bold=True)
add_para("START \u2192 术语归一化 \u2192 条件分支：")
add_para("  \u251c matched     \u2192 SQL拼装 \u2192 查视图 \u2192 格式化 \u2192 END")
add_para("  \u251c ambiguous   \u2192 追问消息 \u2192 END（用户选择后重入）")
add_para("  \u251c multi_match \u2192 逐指标查询 \u2192 合并输出 \u2192 END")
add_para("  \u2514 no_match    \u2192 现有NL2SQL工作流 \u2192 END")

# ═══════════════════════════════════════════
# 7. 实施路线图
# ═══════════════════════════════════════════
add_heading("七、实施路线图", 1)

add_table(
    ["阶段", "时间", "内容", "产出"],
    [
        ["阶段0：边界确认", "第1~2周", "访谈核心用户，收集高频问题", "项目边界确认书（三方签字）"],
        ["阶段1：元数据标注", "第3~8周", "标注核心表结构与字段口径", "字段级口径定义文档"],
        ["阶段2：术语标准化", "第3~6周", "统一业务术语定义", "MES业务术语标准手册"],
        ["阶段3：指标平台", "第9~16周", "建17个指标视图，业务方确认", "视图部署到生产库只读Schema"],
        ["阶段4：路由层接入", "第17~19周", "实现术语归一化 + 路由分发", "Dify工作流上线灰度"],
        ["持续运营", "第20周起", "每周审计未匹配日志，扩充字典", "指标覆盖持续扩大"],
    ],
)

# ═══════════════════════════════════════════
# 8. 前置条件
# ═══════════════════════════════════════════
add_heading("八、前置条件（上线前必须完成）", 1)

add_para(f"1. 确定第一批用户是谁（具体姓名和岗位），不接受{LQ}所有人{RQ}这种答案。")
add_para("2. 收集他们最频繁的 10 个查询问题 \u2192 决定第一期指标范围。")
add_para(f"3. 确认{LQ}良品率{RQ}的计算口径由谁定义（生产部 vs 质量部）。")
add_para("4. 确认跨系统编码是否一致（MES 的 pdline_code 和 ERP 的产线编码是同一套吗？）")
add_para(f"5. 产出《MES 业务术语标准手册》（{LQ}完工{RQ}{LQ}在制{RQ}{LQ}可用库存{RQ}的精确定义）。")
add_para("6. 4 个待确认视图的字段必须先向业务方核实（M001/M006/M014/M016/M017）。")

# ═══════════════════════════════════════════
# 9. 关键决策
# ═══════════════════════════════════════════
add_heading("九、关键决策", 1)

add_para("1. 视图必须预先创建，不能用时创建。用时创建等于把口径确认推迟到查询发生时，出了问题没人负责。", bold=True)
add_para("2. 术语不统一时，指标平台做别名是近期手段，强制统一术语是长期目标。两件事都要做。", bold=True)
add_para("3. 别名字典放代码节点（精确匹配），不放知识库（向量检索）。", bold=True)
add_para("4. 歧义术语宁可追问一次，不要少确认一次。", bold=True)
add_para("5. 先从无争议指标开始（如在制工单数），让业务方先看到价值，再推进敏感指标。", bold=True)

# ── 保存 ──
output_path = "docs/MES智能问数技术方案_简要说明.docx"
doc.save(output_path)
print(f"已生成: {output_path}")
